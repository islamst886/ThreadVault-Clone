"""
docx_generator.py
-----------------
Generates a formatted .docx file from a list of extracted Reddit PostData objects.

Document structure:
  ┌─ Page 1: ThreadVault Research Summary
  │   ├─ Category Breakdown table
  │   ├─ Top Pain Points Discovered
  │   ├─ Products & Competitors Mentioned
  │   ├─ Opportunity Signals (Solution Requests)
  │   └─ Willingness to Pay
  │   [PAGE BREAK]
  ├─ Document Title (ThreadVault Export — query — date)
  ├─ ═══ POST 1 of N ═══
  │   ├─ Post Title (Heading 2 style, bold, 14pt)
  │   ├─ Metadata line (subreddit, author, date, score, comments)
  │   ├─ [AI Analysis block — shaded gray box]
  │   ├─ Post body text
  │   ├─ [Media note if applicable]
  │   └─ COMMENTS
  │       ├─ u/author | ⬆ score | date
  │       │   Comment body
  │       └─   ↳ u/reply_author | ⬆ score | date  (indented per depth)
  │              Reply body
  └─ [page break] POST 2 of N ...
"""

import logging
import os
import re
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Logging ────────────────────────────────────────────────────────────────────
logger = logging.getLogger(__name__)
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)

# ── Design constants ───────────────────────────────────────────────────────────
FONT_BODY     = "Calibri"
FONT_HEADING  = "Calibri"

SIZE_DOC_TITLE   = Pt(18)
SIZE_POST_TITLE  = Pt(14)
SIZE_METADATA    = Pt(10)
SIZE_BODY        = Pt(11)
SIZE_COMMENT_HDR = Pt(10)
SIZE_COMMENT_BOD = Pt(10)
SIZE_SECTION_HDR = Pt(11)
SIZE_SUMMARY_H1  = Pt(20)
SIZE_SUMMARY_H2  = Pt(13)
SIZE_AI_LABEL    = Pt(9)
SIZE_AI_BODY     = Pt(10)

# Colors (RGB tuples → RGBColor)
COLOR_BLACK      = RGBColor(0x1A, 0x1A, 0x1A)   # Near-black for body text
COLOR_GRAY       = RGBColor(0x88, 0x88, 0x88)   # Muted metadata / media notes
COLOR_DARK_GRAY  = RGBColor(0x44, 0x44, 0x44)   # Comment headers
COLOR_ACCENT     = RGBColor(0x1A, 0x59, 0xB1)   # ThreadVault blue for doc title
COLOR_SEPARATOR  = RGBColor(0xCC, 0xCC, 0xCC)   # Horizontal rule tint
COLOR_AI_LABEL   = RGBColor(0x44, 0x44, 0x66)   # Purple-grey for AI field labels
COLOR_GOLD       = RGBColor(0xB4, 0x5A, 0x00)   # Amber for opportunity / WTP
COLOR_GREEN      = RGBColor(0x0A, 0x7A, 0x3A)   # Green for positive WTP signal
COLOR_RED        = RGBColor(0xAA, 0x22, 0x22)   # Red for negative sentiment
COLOR_SHD_GRAY   = "F0F0F4"                       # Light lavender-grey shading hex

# Indentation per comment depth level (in inches)
INDENT_PER_DEPTH = 0.4

# Max items shown in summary lists (0 = unlimited)
SUMMARY_PAIN_TOP   = 10
SUMMARY_PRODUCT_MAX = 0   # all
SUMMARY_SOLUTION_MAX = 0  # all


# ── Text sanitisation ──────────────────────────────────────────────────────────

_ILLEGAL_XML_CHARS_RE = re.compile(
    r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x84\x86-\x9f\ud800-\udfff\ufffe\uffff]"
)

def _sanitize_data(data: Any) -> Any:
    """Recursively strip invalid XML characters from all strings in the data structure."""
    if isinstance(data, dict):
        return {k: _sanitize_data(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [_sanitize_data(v) for v in data]
    elif isinstance(data, str):
        return _ILLEGAL_XML_CHARS_RE.sub("", data)
    return data


# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _add_horizontal_rule(doc: Document, color_hex: str = "BBBBBB", thickness: int = 6) -> None:
    """
    Inserts a horizontal rule (paragraph bottom border) into the document.

    python-docx has no native HR support, so we inject it via raw OOXML.

    Args:
        doc:        Target Document object.
        color_hex:  6-character hex color string (without '#').
        thickness:  Border thickness in eighths of a point.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)

    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_paragraph_indent(paragraph, depth: int) -> None:
    """
    Sets the left indent of a paragraph based on nesting depth.

    Args:
        paragraph: python-docx Paragraph object.
        depth:     Nesting level (0 = no indent, 1 = one level in, etc.).
    """
    paragraph.paragraph_format.left_indent = Inches(INDENT_PER_DEPTH * depth)


def _set_paragraph_shading(paragraph, fill_hex: str = COLOR_SHD_GRAY) -> None:
    """
    Applies a background fill shading to a paragraph via raw OOXML.

    Args:
        paragraph: python-docx Paragraph object.
        fill_hex:  6-character hex fill color (without '#').
    """
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _set_left_border(paragraph, color_hex: str = "4444AA", thickness: int = 18) -> None:
    """
    Adds a thick left border to a paragraph, creating a 'callout' effect.

    Args:
        paragraph:  python-docx Paragraph object.
        color_hex:  6-character hex border color.
        thickness:  Border thickness in eighths of a point.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr_list = pPr.findall(qn("w:pBdr"))
    if pBdr_list:
        pBdr = pBdr_list[0]
    else:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(thickness))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)


# ── Run-level formatting helpers ───────────────────────────────────────────────

def _apply_run_style(
    run,
    size: Pt,
    bold: bool = False,
    italic: bool = False,
    color: Optional[RGBColor] = None,
    font_name: str = FONT_BODY,
) -> None:
    """
    Applies font-level formatting to a single run in one call.

    Args:
        run:       python-docx Run object.
        size:      Font size (use Pt(n)).
        bold:      Whether to bold the run.
        italic:    Whether to italicise the run.
        color:     RGBColor instance, or None to keep default.
        font_name: Font family name string.
    """
    run.font.name   = font_name
    run.font.size   = size
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


# ── AI Analysis aggregation helpers ───────────────────────────────────────────

def _collect_ai_data(extracted_posts: list[dict]) -> dict:
    """
    Aggregates AI analysis fields across all posts for the summary page.

    Returns a dict with:
        category_counts   : Counter {category → count}
        pain_points       : Counter {pain_point → count}
        products          : Counter {product → count}
        solutions         : Counter {solution → count}
        wtp_count         : int  (posts with willingness_to_pay == True)
        analysed_count    : int  (posts that have a valid ai_analysis dict)
    """
    category_counts = Counter()
    pain_points     = Counter()
    products        = Counter()
    solutions       = Counter()
    wtp_count       = 0
    analysed_count  = 0

    for wrapper in extracted_posts:
        ai = wrapper.get("ai_analysis")
        if not isinstance(ai, dict):
            continue
        analysed_count += 1

        cat = (ai.get("category") or "Other").strip()
        category_counts[cat] += 1

        for item in (ai.get("pain_points") or []):
            key = (item or "").strip()
            if key:
                pain_points[key] += 1

        for item in (ai.get("mentioned_products") or []):
            key = (item or "").strip()
            if key:
                products[key] += 1

        for item in (ai.get("solution_requests") or []):
            key = (item or "").strip()
            if key:
                solutions[key] += 1

        if ai.get("willingness_to_pay") is True:
            wtp_count += 1

    return {
        "category_counts":  category_counts,
        "pain_points":      pain_points,
        "products":         products,
        "solutions":        solutions,
        "wtp_count":        wtp_count,
        "analysed_count":   analysed_count,
    }


# ── Summary page section builders ─────────────────────────────────────────────

def _add_summary_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    _apply_run_style(run, SIZE_SUMMARY_H1, bold=True, color=COLOR_ACCENT,
                     font_name=FONT_HEADING)


def _add_summary_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    _apply_run_style(run, SIZE_METADATA, color=COLOR_GRAY)


def _add_summary_section_heading(doc: Document, text: str) -> None:
    """Bolded sub-heading for each summary section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    _apply_run_style(run, SIZE_SUMMARY_H2, bold=True, color=COLOR_BLACK,
                     font_name=FONT_HEADING)
    _add_horizontal_rule(doc, color_hex="CCCCCC", thickness=4)


def _add_summary_numbered_list(
    doc: Document,
    items: list[tuple[str, int]],  # (text, count)
    count_label: str = "posts",
    show_count: bool = True,
) -> None:
    """
    Writes a numbered list of (text, count) tuples into the document.

    Args:
        doc:         Target Document.
        items:       List of (item_text, frequency) pairs, pre-sorted.
        count_label: Singular noun after the number, e.g. 'posts'.
        show_count:  Whether to append ' (N posts)' suffix.
    """
    if not items:
        p = doc.add_paragraph()
        run = p.add_run("• None identified.")
        _apply_run_style(run, SIZE_BODY, italic=True, color=COLOR_GRAY)
        return

    for rank, (text, count) in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.25)

        num_run = p.add_run(f"{rank}. ")
        _apply_run_style(num_run, SIZE_BODY, bold=True, color=COLOR_DARK_GRAY)

        text_run = p.add_run(text)
        _apply_run_style(text_run, SIZE_BODY, color=COLOR_BLACK)

        if show_count and count > 1:
            count_run = p.add_run(f"  ({count} {count_label})")
            _apply_run_style(count_run, SIZE_METADATA, color=COLOR_GRAY)


def _add_category_table(doc: Document, category_counts: Counter) -> None:
    """
    Inserts a simple two-column table: Category | Posts.

    Args:
        doc:             Target Document.
        category_counts: Counter mapping category name → post count.
    """
    if not category_counts:
        p = doc.add_paragraph()
        run = p.add_run("• No categories available.")
        _apply_run_style(run, SIZE_BODY, italic=True, color=COLOR_GRAY)
        return

    # Sort descending by count
    rows = sorted(category_counts.items(), key=lambda x: x[1], reverse=True)

    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for cell, label in zip(hdr.cells, ["Category", "Posts"]):
        cell.text = label
        for run in cell.paragraphs[0].runs:
            _apply_run_style(run, SIZE_BODY, bold=True, color=COLOR_BLACK)
        # Shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "E8ECF8")
        tcPr.append(shd)

    # Data rows
    for i, (cat, count) in enumerate(rows, start=1):
        row = table.rows[i]
        row.cells[0].text = cat
        row.cells[1].text = str(count)
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                _apply_run_style(run, SIZE_BODY, color=COLOR_BLACK)

    doc.add_paragraph()  # breathing room after table


def _add_research_summary_page(
    doc: Document,
    query: str,
    date_str: str,
    ai_data: dict,
    total_posts: int,
) -> None:
    """
    Inserts the full Research Summary as page 1, followed by a page break.

    Args:
        doc:        Target Document.
        query:      Original search query.
        date_str:   Pre-formatted UTC date string.
        ai_data:    Dict from _collect_ai_data().
        total_posts: Total number of posts in the export.
    """
    # ── Title & subtitle ──────────────────────────────────────────────────────
    _add_summary_h1(doc, "ThreadVault Research Summary")
    _add_summary_subtitle(
        doc,
        f"Query: \"{query}\"  |  Generated: {date_str}  |  Posts: {total_posts}",
    )
    _add_horizontal_rule(doc, color_hex="1A59B1", thickness=12)

    analysed = ai_data["analysed_count"]
    if analysed == 0:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        run = p.add_run(
            "No AI analysis data is available for this export. "
            "Check that GEMINI_API_KEY is set in your .env file."
        )
        _apply_run_style(run, SIZE_BODY, italic=True, color=COLOR_GRAY)
        doc.add_page_break()
        return

    # ── 1: Category Breakdown ─────────────────────────────────────────────────
    _add_summary_section_heading(doc, "CATEGORY BREAKDOWN")
    _add_category_table(doc, ai_data["category_counts"])

    # ── 2: Top Pain Points ────────────────────────────────────────────────────
    _add_summary_section_heading(doc, "TOP PAIN POINTS DISCOVERED")
    pain_items = ai_data["pain_points"].most_common(
        SUMMARY_PAIN_TOP if SUMMARY_PAIN_TOP > 0 else None
    )
    _add_summary_numbered_list(doc, pain_items, count_label="posts")

    # ── 3: Products & Competitors ─────────────────────────────────────────────
    _add_summary_section_heading(doc, "PRODUCTS & COMPETITORS MENTIONED")
    product_items = ai_data["products"].most_common(
        SUMMARY_PRODUCT_MAX if SUMMARY_PRODUCT_MAX > 0 else None
    )
    _add_summary_numbered_list(doc, product_items, count_label="posts")

    # ── 4: Opportunity Signals ────────────────────────────────────────────────
    _add_summary_section_heading(doc, "OPPORTUNITY SIGNALS  (Solution Requests)")
    solution_items = ai_data["solutions"].most_common(
        SUMMARY_SOLUTION_MAX if SUMMARY_SOLUTION_MAX > 0 else None
    )
    _add_summary_numbered_list(doc, solution_items, count_label="posts")

    # ── 5: Willingness to Pay ─────────────────────────────────────────────────
    _add_summary_section_heading(doc, "WILLINGNESS TO PAY")
    wtp   = ai_data["wtp_count"]
    total = ai_data["analysed_count"]
    pct   = round(wtp / total * 100) if total else 0
    strong_signal = pct > 20

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    wtp_color = COLOR_GREEN if strong_signal else COLOR_DARK_GRAY
    wtp_run = p.add_run(f"{wtp} out of {total} posts ({pct}%) contain pricing or budget discussion.")
    _apply_run_style(wtp_run, SIZE_BODY, bold=strong_signal, color=wtp_color)

    if strong_signal:
        p2 = doc.add_paragraph()
        signal_run = p2.add_run(
            "✅ Strong market validation signal — significant portion of the audience "
            "is already thinking about price."
        )
        _apply_run_style(signal_run, SIZE_BODY, italic=True, color=COLOR_GREEN)

    # ── Page break before the main post content ───────────────────────────────
    doc.add_page_break()


# ── Per-post AI analysis block ─────────────────────────────────────────────────

def _add_ai_analysis_block(doc: Document, ai: dict) -> None:
    """
    Inserts a shaded AI analysis block after the post metadata line.

    The block contains: Category, Sentiment, Summary, Key Quote,
    Pain Points, and Solutions Requested — all in a light gray shaded box
    with a purple-blue left border.

    Args:
        doc: Target Document.
        ai:  The ``ai_analysis`` dict from the post wrapper (may be partial).
    """
    if not isinstance(ai, dict):
        return

    def _shaded_para(indent_in: float = 0.2) -> Any:
        """Creates a shaded, bordered paragraph and returns it."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(indent_in)
        p.paragraph_format.right_indent = Inches(0.1)
        _set_paragraph_shading(p, fill_hex=COLOR_SHD_GRAY)
        _set_left_border(p, color_hex="4444AA", thickness=16)
        return p

    def _label_run(p, label: str) -> None:
        run = p.add_run(label)
        _apply_run_style(run, SIZE_AI_LABEL, bold=True, color=COLOR_AI_LABEL)

    def _value_run(p, value: str, italic: bool = False,
                   color: RGBColor = COLOR_BLACK) -> None:
        run = p.add_run(value)
        _apply_run_style(run, SIZE_AI_BODY, italic=italic, color=color)

    # ── Spacer line at top of block ──────────────────────────────────────────
    top_spacer = doc.add_paragraph()
    top_spacer.paragraph_format.space_before = Pt(4)
    top_spacer.paragraph_format.space_after  = Pt(0)
    _set_paragraph_shading(top_spacer)
    _set_left_border(top_spacer, color_hex="4444AA", thickness=16)

    # ── Row 1: Category + Sentiment on same line ─────────────────────────────
    category  = ai.get("category",  "—") or "—"
    sentiment = (ai.get("sentiment", "neutral") or "neutral").lower()

    sent_color_map = {
        "positive": COLOR_GREEN,
        "negative": COLOR_RED,
        "neutral":  COLOR_DARK_GRAY,
        "mixed":    COLOR_GOLD,
    }
    sent_color = sent_color_map.get(sentiment, COLOR_DARK_GRAY)

    p1 = _shaded_para()
    _label_run(p1, "Category: ")
    _value_run(p1, category, color=COLOR_ACCENT)
    _label_run(p1, "    Sentiment: ")
    _value_run(p1, sentiment.capitalize(), color=sent_color)

    # ── Row 2: Summary ───────────────────────────────────────────────────────
    summary = (ai.get("summary") or "").strip()
    if summary:
        p2 = _shaded_para()
        _label_run(p2, "Summary: ")
        _value_run(p2, summary)

    # ── Row 3: Key Quote ─────────────────────────────────────────────────────
    key_quote = (ai.get("key_quote") or "").strip()
    if key_quote:
        p3 = _shaded_para()
        _label_run(p3, "Key Quote: ")
        _value_run(p3, f'"{key_quote}"', italic=True, color=COLOR_DARK_GRAY)

    # ── Row 4: Pain Points ───────────────────────────────────────────────────
    pain_points = [s.strip() for s in (ai.get("pain_points") or []) if s and s.strip()]
    if pain_points:
        pp = _shaded_para()
        _label_run(pp, "Pain Points: ")

        for item in pain_points:
            pi = _shaded_para(indent_in=0.45)
            _value_run(pi, f"• {item}")

    # ── Row 5: Solutions Requested ───────────────────────────────────────────
    solutions = [s.strip() for s in (ai.get("solution_requests") or []) if s and s.strip()]
    if solutions:
        sp = _shaded_para()
        _label_run(sp, "Solutions Requested: ")

        for item in solutions:
            si = _shaded_para(indent_in=0.45)
            _value_run(si, f"• {item}", color=COLOR_GOLD)

    # ── Bottom spacer ────────────────────────────────────────────────────────
    bot_spacer = doc.add_paragraph()
    bot_spacer.paragraph_format.space_before = Pt(0)
    bot_spacer.paragraph_format.space_after  = Pt(6)
    _set_paragraph_shading(bot_spacer)
    _set_left_border(bot_spacer, color_hex="4444AA", thickness=16)


# ── Document-level section builders ───────────────────────────────────────────

def _add_document_title(doc: Document, query: str, date_str: str) -> None:
    """
    Inserts the top-of-document title heading.

    Format: "ThreadVault Export — [query] — [date]"

    Args:
        doc:      Target Document.
        query:    The original user search query.
        date_str: Formatted date string (e.g. "2024-12-01 14:30 UTC").
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    # "ThreadVault Export" in accent color
    run_brand = p.add_run("ThreadVault Export")
    _apply_run_style(run_brand, SIZE_DOC_TITLE, bold=True, color=COLOR_ACCENT)

    run_sep = p.add_run(" — ")
    _apply_run_style(run_sep, SIZE_DOC_TITLE, color=COLOR_GRAY)

    run_query = p.add_run(query)
    _apply_run_style(run_query, SIZE_DOC_TITLE, bold=True, color=COLOR_BLACK)

    run_date = p.add_run(f" — {date_str}")
    _apply_run_style(run_date, SIZE_METADATA, color=COLOR_GRAY)

    _add_horizontal_rule(doc, color_hex="1A59B1", thickness=12)


def _add_post_separator(doc: Document, post_num: int, total: int) -> None:
    """
    Inserts a visual separator between posts: "═══ POST X of Y ═══"

    Args:
        doc:      Target Document.
        post_num: 1-based index of this post.
        total:    Total number of posts being exported.
    """
    _add_horizontal_rule(doc, color_hex="BBBBBB", thickness=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)

    run = p.add_run(f"POST {post_num} of {total}")
    _apply_run_style(run, SIZE_SECTION_HDR, bold=True, color=COLOR_GRAY)


def _add_post_title(doc: Document, title: str) -> None:
    """
    Inserts the post title as a prominent heading.

    Args:
        doc:   Target Document.
        title: Reddit post title string.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)

    run = p.add_run(title)
    _apply_run_style(run, SIZE_POST_TITLE, bold=True, color=COLOR_BLACK,
                     font_name=FONT_HEADING)


def _add_post_metadata(doc: Document, post: dict, sort: str = None, limit: Any = None) -> None:
    """
    Inserts the metadata line below the post title.

    Args:
        doc:   Target Document.
        post:   Post data dictionary.
        sort:   Sort mode used.
        limit:  Comment limit used.
    """
    flair_part = f" | 🏷 {post.get('flair')}" if post.get("flair") else ""
    sort_info = f"  |  🔍 Sort: {str(sort).capitalize()}" if sort else ""
    limit_info = f"  |  📊 Limit: {str(limit).capitalize()}" if limit else ""

    meta_text = (
        f"r/{post.get('subreddit')}  |  Posted by u/{post.get('author')}  |  {post.get('posted_at')}"
        f"  |  ⬆ {post.get('score', 0):,} upvotes ({post.get('upvote_ratio', '0%')} upvoted)"
        f"  |  💬 {post.get('num_comments', 0):,} comments{flair_part}{sort_info}{limit_info}"
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(meta_text)
    _apply_run_style(run, SIZE_METADATA, color=COLOR_GRAY)


def _add_post_body(doc: Document, post: dict) -> None:
    """
    Inserts the post body text, or the media note if no text content.

    Args:
        doc:  Target Document.
        post: PostData instance.
    """
    if post.get("media_note"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(post["media_note"])
        _apply_run_style(run, SIZE_BODY, italic=True, color=COLOR_GRAY)

    body = post.get("body", "")
    if body and body not in ("[removed]", "[deleted]", ""):
        # Split body on double newlines to preserve paragraph breaks
        paragraphs = body.split("\n\n")
        for para_text in paragraphs:
            text = para_text.strip()
            if not text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text.replace("\n", " "))
            _apply_run_style(run, SIZE_BODY, color=COLOR_BLACK)

    elif body in ("[removed]", "[deleted]"):
        p = doc.add_paragraph()
        run = p.add_run(body)
        _apply_run_style(run, SIZE_BODY, italic=True, color=COLOR_GRAY)


# ── Comment section builders ───────────────────────────────────────────────────

def _add_comment_block(doc: Document, comment: dict) -> None:
    """
    Recursively writes a comment and all its replies into the document.

    Indentation increases with depth using left_indent.
    Replies are prefixed with "↳" to show nesting visually.

    Args:
        doc:     Target Document.
        comment: CommentData instance (may contain nested replies).
    """
    depth = comment.get("depth", 0)

    # ── Comment header line: u/author | ⬆ score | date ───────────────────────
    header_p = doc.add_paragraph()
    _set_paragraph_indent(header_p, depth)
    header_p.paragraph_format.space_before = Pt(6) if depth == 0 else Pt(3)
    header_p.paragraph_format.space_after  = Pt(1)

    if depth > 0:
        arrow = header_p.add_run("↳ ")
        _apply_run_style(arrow, SIZE_COMMENT_HDR, color=COLOR_GRAY)

    author_run = header_p.add_run(f"u/{comment.get('author')}")
    _apply_run_style(author_run, SIZE_COMMENT_HDR, bold=True, color=COLOR_DARK_GRAY)

    meta_run = header_p.add_run(
        f"  |  ⬆ {comment.get('score', 0):,}  |  {comment.get('posted_at', '')}"
    )
    _apply_run_style(meta_run, SIZE_COMMENT_HDR, color=COLOR_GRAY)

    # ── Comment body ──────────────────────────────────────────────────────────
    if comment.get("media_note"):
        media_p = doc.add_paragraph()
        _set_paragraph_indent(media_p, depth)
        media_p.paragraph_format.space_after = Pt(2)
        media_run = media_p.add_run(comment["media_note"])
        _apply_run_style(media_run, SIZE_COMMENT_BOD, italic=True, color=COLOR_GRAY)

    body = comment.get("body", "")
    if body and body not in ("[removed]", "[deleted]", ""):
        body_text = body.replace("\n", " ").strip()
        body_p = doc.add_paragraph()
        _set_paragraph_indent(body_p, depth)
        body_p.paragraph_format.space_after = Pt(2)
        body_run = body_p.add_run(body_text)
        _apply_run_style(body_run, SIZE_COMMENT_BOD, color=COLOR_BLACK)
    elif body in ("[removed]", "[deleted]"):
        body_p = doc.add_paragraph()
        _set_paragraph_indent(body_p, depth)
        body_run = body_p.add_run(body)
        _apply_run_style(body_run, SIZE_COMMENT_BOD, italic=True, color=COLOR_GRAY)

    # ── Recurse into replies ──────────────────────────────────────────────────
    for reply in comment.get("replies", []):
        _add_comment_block(doc, reply)


def _add_comments_section(doc: Document, comments: list[dict]) -> None:
    """
    Inserts the "COMMENTS" section header followed by all top-level comments.

    Args:
        doc:      Target Document.
        comments: List of top-level CommentData objects.
    """
    if not comments:
        p = doc.add_paragraph()
        run = p.add_run("No comments extracted.")
        _apply_run_style(run, SIZE_METADATA, italic=True, color=COLOR_GRAY)
        return

    # Section header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"COMMENTS  ({len(comments):,} top-level)")
    _apply_run_style(run, SIZE_SECTION_HDR, bold=True, color=COLOR_BLACK)

    # Thin rule under "COMMENTS"
    _add_horizontal_rule(doc, color_hex="DDDDDD", thickness=4)

    for comment in comments:
        _add_comment_block(doc, comment)


# ── Main generation entry point ────────────────────────────────────────────────

def generate_docx(
    extracted_posts: list[dict],
    query: str,
    output_dir: str = "output",
    progress_callback: Optional[Callable[[int, int], None]] = None,
) -> str:
    """
    Generates a formatted .docx file from a list of dict objects.

    Args:
        extracted_posts: List of post dictionary objects.
        query:           The original search query entered by the user.
        output_dir:      Directory where the .docx will be saved. Created if missing.
        progress_callback: Optional callable(idx, total) for live progress updates.

    Returns:
        Absolute file path of the generated .docx file.

    Raises:
        ValueError: If `posts` is empty.
        OSError:    If the output directory cannot be created.
    """
    if not extracted_posts:
        raise ValueError("No posts provided — nothing to export.")

    # Sanitise all incoming data to prevent python-docx XML export crashes
    extracted_posts = _sanitize_data(extracted_posts)
    query = _sanitize_data(query)

    # ── Prepare output directory ───────────────────────────────────────────────
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    # ── Timestamp + query slug for a human-readable filename ──────────────────
    now_utc    = datetime.now(tz=timezone.utc)
    date_str   = now_utc.strftime("%Y-%m-%d %H:%M UTC")
    file_stamp = now_utc.strftime("%Y%m%d_%H%M%S")

    # Build a safe slug from the query:
    #   - lowercase, replace any run of non-alphanumeric chars with _
    #   - strip leading/trailing underscores, cap at 40 chars
    query_slug = re.sub(r"[^a-z0-9]+", "_", query.lower().strip()).strip("_")[:40]
    post_count = len(extracted_posts)
    filename   = f"{query_slug}_{post_count}posts_{file_stamp}.docx"
    filepath   = os.path.join(output_dir, filename)

    logger.info(
        "Generating DOCX for query '%s' with %d posts → %s",
        query, len(extracted_posts), filepath,
    )

    # ── Create document and set default style ─────────────────────────────────
    doc = Document()

    # Set document-wide default font via the Normal style
    normal_style = doc.styles["Normal"]
    normal_style.font.name = FONT_BODY
    normal_style.font.size = SIZE_BODY

    # ── Set narrow-ish margins ─────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Page 1: Research Summary ───────────────────────────────────────────────
    ai_data = _collect_ai_data(extracted_posts)
    _add_research_summary_page(doc, query, date_str, ai_data, total_posts=post_count)

    # ── Document title (after summary page) ───────────────────────────────────
    _add_document_title(doc, query, date_str)

    total = len(extracted_posts)

    # ── Iterate posts ──────────────────────────────────────────────────────────
    for idx, wrapper in enumerate(extracted_posts, start=1):
        if progress_callback:
            progress_callback(idx, total)

        post     = wrapper["post"]
        comments = wrapper["comments"]
        ai       = wrapper.get("ai_analysis")
        title    = post.get("title", "(No Title)")
        logger.info("Writing post %d / %d: %s", idx, total, title[:60])

        _add_post_separator(doc, idx, total)
        _add_post_title(doc, title)
        _add_post_metadata(doc, post, wrapper.get("sort_used"), wrapper.get("limit_used"))

        # AI block — inserted immediately after metadata, before body
        if isinstance(ai, dict):
            _add_ai_analysis_block(doc, ai)

        _add_post_body(doc, post)
        _add_comments_section(doc, comments)

        # Page break after every post except the last
        if idx < total:
            doc.add_page_break()

    # ── Save ───────────────────────────────────────────────────────────────────
    doc.save(filepath)
    abs_path = os.path.abspath(filepath)
    logger.info("DOCX saved: %s", abs_path)
    return abs_path


# ══════════════════════════════════════════════════════════════════════════════
# BULK DOCX GENERATION — generate_bulk_docx() and its private helpers
# ══════════════════════════════════════════════════════════════════════════════

_BULK_CHECKPOINT = 50   # save to disk every N posts (crash protection)
_MAX_DESC_CHARS  = 200  # truncate subreddit descriptions


def _fmt_members(n: int) -> str:
    """Formats subscriber count: 234000 → '234K', 12100000 → '12.1M'."""
    if n >= 1_000_000:
        return f"{n / 1_000_000:.1f}M"
    if n >= 1_000:
        return f"{n / 1_000:.0f}K"
    return str(n)


def _extract_year(posted_at: str) -> str:
    """Extracts the 4-digit year from 'YYYY-MM-DD HH:MM UTC'."""
    try:
        return posted_at[:4]
    except (IndexError, TypeError):
        return "?"


def _safe_slug(name: str, max_len: int = 30) -> str:
    """Returns a filesystem-safe lowercase slug."""
    import re as _re
    return _re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")[:max_len]


def _add_bulk_hyperlink(paragraph, display: str, url: str) -> None:
    """Inserts a clickable hyperlink run into *paragraph* via raw OOXML."""
    from docx.oxml import OxmlElement as _oel
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = _oel("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    hl.set(qn("w:history"), "1")
    run = _oel("w:r")
    rPr = _oel("w:rPr")
    rSty = _oel("w:rStyle")
    rSty.set(qn("w:val"), "Hyperlink")
    rPr.append(rSty)
    run.append(rPr)
    t = _oel("w:t")
    t.text = display
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)


# ── Bulk: summary page (Page 1 of each file) ──────────────────────────────────

def _build_bulk_summary_page(
    doc: "Document",
    file_subs: list[dict],
    meta: dict,
    now_utc: datetime,
    file_label: str = "",
) -> None:
    """
    Writes the master summary page for one bulk DOCX file then inserts a page
    break.  *file_subs* is the subset of subreddits included in THIS file.
    """
    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(4)
    _apply_run_style(
        tp.add_run("ThreadVault Research Export"),
        SIZE_SUMMARY_H1, bold=True, color=COLOR_ACCENT, font_name=FONT_HEADING,
    )

    # Subtitles
    total_posts = sum(len(si.get("posts", [])) for si in file_subs)
    yr = meta.get("years_back", 1)
    try:
        yr_int = int(yr)
    except:
        yr_int = 1
        
    sub_names = [f"r/{si.get('name', 'sub')}" for si in file_subs]
    sub_list_str = " + ".join(sub_names) if len(sub_names) <= 3 else f"{len(sub_names)} Subreddits"
    
    line1_parts = [sub_list_str, f"{total_posts} posts", f"Past {yr_int} year{'s' if yr_int > 1 else ''}"]
    if file_label:
        line1_parts.append(file_label)
        
    sp1 = doc.add_paragraph()
    sp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp1.paragraph_format.space_after = Pt(2)
    _apply_run_style(sp1.add_run("  •  ".join(line1_parts)), SIZE_METADATA, color=COLOR_GRAY)

    date_str = now_utc.strftime("%d %B %Y").lstrip("0")
    sp2 = doc.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp2.paragraph_format.space_after = Pt(12)
    _apply_run_style(sp2.add_run(f"Generated: {date_str}"), SIZE_METADATA, color=COLOR_GRAY)

    _add_horizontal_rule(doc, color_hex="1A59B1", thickness=12)

    # Section heading
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(10)
    hp.paragraph_format.space_after  = Pt(6)
    _apply_run_style(hp.add_run("SUBREDDITS IN THIS FILE"), SIZE_SUMMARY_H2,
                     bold=True, color=COLOR_BLACK, font_name=FONT_HEADING)

    # Table: Subreddit | Members | Posts | Date Range
    tbl = doc.add_table(rows=1 + len(file_subs), cols=4)
    tbl.style = "Table Grid"
    for cell, label in zip(tbl.rows[0].cells, ["Subreddit", "Members", "Posts", "Date Range"]):
        cell.text = label
        for run in cell.paragraphs[0].runs:
            _apply_run_style(run, SIZE_BODY, bold=True, color=COLOR_BLACK)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "E8ECF8")
        tcPr.append(shd)

    total_posts = total_comments = 0
    skipped_notes: list[str] = []

    for i, si in enumerate(file_subs, start=1):
        about      = si.get("about", {})
        posts      = si.get("posts", [])
        name       = si.get("name", "?")
        total_fnd  = si.get("total_found", len(posts))
        limit_hit  = si.get("limit_hit", False)
        total_posts    += len(posts)
        total_comments += sum(p.get("total_comments_extracted", 0) for p in posts)

        years = [_extract_year(p.get("post", {}).get("posted_at", "")) for p in posts]
        years = [y for y in years if y != "?"]
        date_range = f"{min(years)} – {max(years)}" if years else "—"

        row = tbl.rows[i]
        row.cells[0].text = f"r/{name}"
        row.cells[1].text = _fmt_members(about.get("subscribers", 0))
        row.cells[2].text = str(len(posts))
        row.cells[3].text = date_range
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                _apply_run_style(run, SIZE_BODY, color=COLOR_BLACK)

        if limit_hit and total_fnd > len(posts):
            skipped_notes.append(
                f"⚠  r/{name} had {total_fnd:,} posts in this period. "
                f"Only the first {len(posts):,} were extracted due to the post limit setting. "
                "To get more, increase the limit in the extraction settings."
            )

    doc.add_paragraph()  # breathing room

    # Totals
    def _stat_line(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _apply_run_style(p.add_run(f"{label}: "), SIZE_BODY, bold=True, color=COLOR_DARK_GRAY)
        _apply_run_style(p.add_run(value), SIZE_BODY, color=COLOR_BLACK)

    _stat_line("Total posts in this file", f"{total_posts:,}")
    _stat_line("Total comments in this file", f"{total_comments:,}")

    sort_label  = str(meta.get("comment_sort", "top")).capitalize()
    limit_label = str(meta.get("comment_limit", 25)).capitalize()
    _stat_line("Extraction settings", f"Sort: {sort_label}  |  Comments per post: {limit_label}")

    # Skipped post warnings
    if skipped_notes:
        doc.add_paragraph()
        for note in skipped_notes:
            np = doc.add_paragraph()
            np.paragraph_format.space_before = Pt(4)
            np.paragraph_format.space_after  = Pt(4)
            _set_paragraph_shading(np, fill_hex="FFF8E1")
            _apply_run_style(np.add_run(note), SIZE_BODY, italic=True, color=COLOR_GOLD)

    doc.add_page_break()


# ── Bulk: subreddit section header ────────────────────────────────────────────

def _add_bulk_subreddit_header(doc: "Document", si: dict) -> None:
    """
    Inserts a subreddit header page before that subreddit's posts, then a page
    break.  *si* is the subreddit info dict (name, about, posts, …).
    """
    about      = si.get("about", {})
    name       = si.get("name", "?")
    posts      = si.get("posts", [])
    total_fnd  = si.get("total_found", len(posts))

    # r/subredditname heading
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(6)
    hp.paragraph_format.space_after  = Pt(4)
    _apply_run_style(hp.add_run(f"r/{name}"), SIZE_SUMMARY_H1, bold=True,
                     color=COLOR_ACCENT, font_name=FONT_HEADING)
    _add_horizontal_rule(doc, color_hex="1A59B1", thickness=10)

    def _field(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _apply_run_style(p.add_run(f"{label}: "), SIZE_BODY, bold=True, color=COLOR_DARK_GRAY)
        _apply_run_style(p.add_run(value), SIZE_BODY, color=COLOR_BLACK)

    full_title = (about.get("title") or "").strip()
    if full_title:
        _field("Title", full_title)

    desc = (about.get("public_description") or about.get("description") or "").strip()
    if desc:
        _field("Description", desc[:_MAX_DESC_CHARS] + ("…" if len(desc) > _MAX_DESC_CHARS else ""))

    _field("Members", _fmt_members(about.get("subscribers", 0)))
    _field("Posts extracted", f"{len(posts):,} of {total_fnd:,} in this period")

    years = [_extract_year(p.get("post", {}).get("posted_at", "")) for p in posts]
    years = [y for y in years if y != "?"]
    if years:
        _field("Date range", f"{min(years)} to {max(years)}")

    doc.add_page_break()


# ── Bulk: comment renderer ────────────────────────────────────────────────────

def _add_bulk_comment_block(doc: "Document", comment: dict) -> None:
    """Recursively writes one comment and all its replies."""
    depth = comment.get("depth", 0)
    body  = comment.get("body", "")
    indent = Inches(min(INDENT_PER_DEPTH * depth, 3.0))

    # Detect "more replies" stub
    is_stub = body.startswith("[") and "more replies not loaded" in body

    if is_stub:
        sp = doc.add_paragraph()
        sp.paragraph_format.left_indent  = indent
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after  = Pt(2)
        _apply_run_style(sp.add_run(body), SIZE_COMMENT_BOD, italic=True, color=COLOR_GRAY)
        return

    # Comment header: → → u/author • ⬆ score • date
    arrows = "→ " * depth
    author = comment.get("author") or "[deleted]"
    score  = comment.get("score", 0)
    date   = comment.get("posted_at", "")

    hdr = doc.add_paragraph()
    hdr.paragraph_format.left_indent  = indent
    hdr.paragraph_format.space_before = Pt(5) if depth == 0 else Pt(2)
    hdr.paragraph_format.space_after  = Pt(1)
    if arrows:
        _apply_run_style(hdr.add_run(arrows), SIZE_COMMENT_HDR, color=COLOR_GRAY)
    _apply_run_style(hdr.add_run(f"u/{author}"), SIZE_COMMENT_HDR, bold=True, color=COLOR_DARK_GRAY)
    _apply_run_style(hdr.add_run(f"  •  ⬆ {score:,}  •  {date}"), SIZE_COMMENT_HDR, color=COLOR_GRAY)

    # Comment body
    if body and body not in ("[removed]", "[deleted]"):
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = indent
        bp.paragraph_format.space_after = Pt(2)
        _apply_run_style(bp.add_run(body.replace("\n", " ").strip()), SIZE_COMMENT_BOD, color=COLOR_BLACK)
    elif body in ("[removed]", "[deleted]"):
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = indent
        _apply_run_style(bp.add_run(body), SIZE_COMMENT_BOD, italic=True, color=COLOR_GRAY)

    # Media note on comment
    if comment.get("media_note"):
        mp = doc.add_paragraph()
        mp.paragraph_format.left_indent = indent
        _apply_run_style(mp.add_run(comment["media_note"]), SIZE_COMMENT_BOD, italic=True, color=COLOR_GRAY)

    for reply in comment.get("replies", []):
        _add_bulk_comment_block(doc, reply)


# ── Bulk: single post renderer ────────────────────────────────────────────────

def _add_bulk_post(
    doc: "Document",
    wrapper: dict,
    post_num: int,
    total_in_sub: int,
    subreddit_name: str,
) -> None:
    """
    Writes one post (header + body + comments) plus a thin separator at the
    bottom.  A page break is inserted every 10 posts (handled by caller).
    """
    post     = wrapper.get("post", {})
    comments = wrapper.get("comments", [])
    limit    = wrapper.get("limit_used", 25)

    title       = post.get("title", "(No Title)")
    author      = post.get("author", "[deleted]")
    score       = post.get("score", 0)
    num_cmts    = post.get("num_comments", 0)
    posted_at   = post.get("posted_at", "")
    subreddit   = post.get("subreddit") or subreddit_name
    url         = post.get("url", "")
    body        = post.get("body", "")
    flair       = post.get("flair")
    media_note  = post.get("media_note")
    total_ext   = wrapper.get("total_comments_extracted", len(comments))

    # ── Post separator banner ──────────────────────────────────────────────────
    _add_horizontal_rule(doc, color_hex="AAAAAA", thickness=4)
    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_p.paragraph_format.space_before = Pt(4)
    sep_p.paragraph_format.space_after  = Pt(2)
    _apply_run_style(
        sep_p.add_run(f"POST {post_num} of {total_in_sub} in r/{subreddit}"),
        SIZE_SECTION_HDR, bold=True, color=COLOR_GRAY,
    )

    # ── Post title ────────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after  = Pt(2)
    _apply_run_style(tp.add_run(title), SIZE_POST_TITLE, bold=True, color=COLOR_BLACK,
                     font_name=FONT_HEADING)

    # ── Metadata line ─────────────────────────────────────────────────────────
    mp = doc.add_paragraph()
    mp.paragraph_format.space_after = Pt(2)
    meta_text = (
        f"r/{subreddit}  •  Posted by u/{author}  •  {posted_at}"
        f"  •  ⬆ {score:,} upvotes  •  {num_cmts:,} comments"
    )
    if url:
        meta_text += f"  •  {url}"
    _apply_run_style(mp.add_run(meta_text), SIZE_METADATA, color=COLOR_GRAY)

    # ── Flair ─────────────────────────────────────────────────────────────────
    if flair:
        fp = doc.add_paragraph()
        fp.paragraph_format.space_after = Pt(2)
        _apply_run_style(fp.add_run(f"[{flair}]"), SIZE_METADATA, color=COLOR_ACCENT)

    # ── Post body ─────────────────────────────────────────────────────────────
    if media_note:
        is_link = "[POST CONTAINS LINK:" in media_note
        if is_link:
            lp = doc.add_paragraph()
            lp.paragraph_format.space_after = Pt(4)
            _apply_run_style(lp.add_run("Link: "), SIZE_BODY, bold=True, color=COLOR_DARK_GRAY)
            _add_bulk_hyperlink(lp, url, url)
        else:
            np = doc.add_paragraph()
            np.paragraph_format.space_after = Pt(4)
            _apply_run_style(np.add_run(media_note), SIZE_BODY, italic=True, color=COLOR_GRAY)

    if body and body not in ("[removed]", "[deleted]"):
        for para_text in body.split("\n\n"):
            text = para_text.strip()
            if not text:
                continue
            bp = doc.add_paragraph()
            bp.paragraph_format.space_after = Pt(4)
            _apply_run_style(bp.add_run(text.replace("\n", " ")), SIZE_BODY, color=COLOR_BLACK)
    elif body in ("[removed]", "[deleted]"):
        rp = doc.add_paragraph()
        _apply_run_style(rp.add_run(body), SIZE_BODY, italic=True, color=COLOR_GRAY)
    elif not media_note:
        ep = doc.add_paragraph()
        _apply_run_style(ep.add_run("[No text content]"), SIZE_BODY, italic=True, color=COLOR_GRAY)

    # ── Comments section ──────────────────────────────────────────────────────
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(8)
    cp.paragraph_format.space_after  = Pt(4)
    _apply_run_style(
        cp.add_run(f"── COMMENTS ({total_ext} shown) ──"),
        SIZE_SECTION_HDR, bold=True, color=COLOR_BLACK,
    )

    if comments:
        for comment in comments:
            _add_bulk_comment_block(doc, comment)
    else:
        np = doc.add_paragraph()
        _apply_run_style(np.add_run("[No comments extracted for this post]"),
                         SIZE_METADATA, italic=True, color=COLOR_GRAY)

    # Bottom thin rule
    _add_horizontal_rule(doc, color_hex="DDDDDD", thickness=3)

# ── Main bulk generation entry point ──────────────────────────────────────────

async def generate_bulk_docx(
    subreddit_data: dict,
    output_dir: str = "output",
    max_posts_per_file: int = 200,
) -> list[str]:
    """
    Generates one or more DOCX files from a multi-subreddit bulk extraction.

    Expected *subreddit_data* schema
    --------------------------------
    {
      "subreddits": [
        {
          "name":        str,        # subreddit name, no "r/"
          "about":       dict,       # Reddit /about.json .data field
          "posts":       list[dict], # list of extract_post() result dicts
          "total_found": int,        # posts found before limit cap
          "limit_hit":   bool,       # True if post_limit was reached
        }, ...
      ],
      "meta": {
        "comment_sort":  str,        # e.g. "top"
        "comment_limit": int | str,  # e.g. 25 or "all"
        "years_back":    int,        # e.g. 2
      }
    }

    File-splitting rules
    --------------------
    Total posts ≤ max_posts_per_file  →  one combined file.
    Total posts  > max_posts_per_file →  one file per subreddit; subreddits
                                         with > max_posts_per_file posts are
                                         further split into numbered parts.

    Returns a list of absolute paths to all generated DOCX files.
    """
    import asyncio as _asyncio

    # ── Sanitise all incoming strings ──────────────────────────────────────────
    subreddit_data = _sanitize_data(subreddit_data)

    subreddits: list[dict] = subreddit_data.get("subreddits", [])
    meta:       dict       = subreddit_data.get("meta", {})

    if not subreddits:
        raise ValueError("subreddit_data contains no subreddits.")

    Path(output_dir).mkdir(parents=True, exist_ok=True)

    now_utc   = datetime.now(tz=timezone.utc)
    ts_date   = now_utc.strftime("%Y-%m-%d")

    total_posts_all = sum(len(si.get("posts", [])) for si in subreddits)
    logger.info(
        "[BulkDOCX] %d subreddit(s), %d total posts, max_per_file=%d",
        len(subreddits), total_posts_all, max_posts_per_file,
    )

    # ── Build the file plan ────────────────────────────────────────────────────
    plan: list[dict] = []

    def _chunks(lst: list, n: int):
        for i in range(0, len(lst), n):
            yield lst[i : i + n]

    yr = meta.get("years_back", 1)

    if total_posts_all <= max_posts_per_file:
        subs_slug = "+".join([f"r-{_safe_slug(si.get('name', 'sub')).replace('_', '-')}" for si in subreddits])
        filename = f"ThreadVault_{subs_slug}_{total_posts_all}posts_{yr}yr_{ts_date}.docx"
        plan.append({
            "filename": filename,
            "subs":     subreddits,
            "label":    "",
        })
    else:
        for si in subreddits:
            name  = si.get("name", "sub")
            posts = si.get("posts", [])
            slug  = _safe_slug(name).replace("_", "-")
            n_posts = len(posts)

            if n_posts <= max_posts_per_file:
                plan.append({
                    "filename": f"ThreadVault_r-{slug}_{n_posts}posts_{yr}yr_{ts_date}.docx",
                    "subs":     [si],
                    "label":    "",
                })
            else:
                chunks = list(_chunks(posts, max_posts_per_file))
                n_parts = len(chunks)
                for part_idx, chunk in enumerate(chunks, start=1):
                    sub_part = {**si, "posts": chunk}
                    plan.append({
                        "filename": f"ThreadVault_r-{slug}_{len(chunk)}posts_{yr}yr_part{part_idx}_{ts_date}.docx",
                        "subs":     [sub_part],
                        "label":    f"Part {part_idx} of {n_parts}",
                    })

    logger.info("[BulkDOCX] File plan: %d file(s).", len(plan))
    generated_paths: list[str] = []

    # ── Generate each planned file ─────────────────────────────────────────────
    for file_spec in plan:
        filename   = file_spec["filename"]
        file_subs  = file_spec["subs"]
        file_label = file_spec.get("label", "")
        filepath   = os.path.join(output_dir, filename)
        abs_path   = os.path.abspath(filepath)

        logger.info("[BulkDOCX] Generating: %s", filename)

        # Create document
        doc = Document()
        normal_style = doc.styles["Normal"]
        normal_style.font.name = FONT_BODY
        normal_style.font.size = SIZE_BODY
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        # Page 1: Master summary
        _build_bulk_summary_page(doc, file_subs, meta, now_utc, file_label)

        # Write posts subreddit by subreddit
        global_post_count = 0   # for checkpoint logic across sub sections

        for si in file_subs:
            sub_name = si.get("name", "unknown")
            posts    = si.get("posts", [])
            n_posts  = len(posts)

            if not posts:
                continue

            # Subreddit header page
            _add_bulk_subreddit_header(doc, si)

            for post_idx, wrapper in enumerate(posts, start=1):
                _add_bulk_post(doc, wrapper, post_idx, n_posts, sub_name)
                global_post_count += 1

                # Page break every 10 posts (except at subreddit boundaries which
                # already have a page break from the subreddit header)
                if post_idx % 10 == 0 and post_idx < n_posts:
                    doc.add_page_break()

                # Checkpoint save every 50 posts
                if global_post_count % _BULK_CHECKPOINT == 0:
                    logger.info(
                        "[BulkDOCX] Checkpoint save at post %d → %s",
                        global_post_count, filename,
                    )
                    await _asyncio.to_thread(doc.save, filepath)

        # Final save
        await _asyncio.to_thread(doc.save, filepath)
        logger.info("[BulkDOCX] Saved: %s", abs_path)
        generated_paths.append(abs_path)

    logger.info(
        "[BulkDOCX] Complete — %d file(s) generated.", len(generated_paths)
    )
    return generated_paths


# ── CLI convenience runner ─────────────────────────────────────────────────────

if __name__ == "__main__":
    """Quick smoke-test with synthetic data — no API calls needed."""
    reply = {
        "author": "replier_user",
        "body": "I totally agree with this point!",
        "score": 45,
        "posted_at": "2024-01-15 09:30 UTC",
        "depth": 1,
        "media_note": None,
        "replies": []
    }
    comment = {
        "author": "top_commenter",
        "body": "This is a really insightful post. Thanks for sharing.",
        "score": 312,
        "posted_at": "2024-01-15 08:00 UTC",
        "depth": 0,
        "media_note": None,
        "replies": [reply]
    }
    post_data = {
        "post": {
            "url": "https://www.reddit.com/r/python/comments/abc123/test/",
            "title": "What are the best Python libraries for data analysis in 2024?",
            "subreddit": "python",
            "author": "demo_user",
            "score": 1842,
            "upvote_ratio": "97%",
            "num_comments": 203,
            "flair": "Discussion",
            "posted_at": "2024-01-15 07:00 UTC",
            "body": (
                "I've been exploring various Python libraries for data analysis and "
                "wanted to get the community's take. Libraries like Pandas, Polars, "
                "and DuckDB all seem promising.\n\n"
                "What has your experience been? Any hidden gems worth trying?"
            ),
            "media_note": None,
        },
        "comments": [comment],
        "sort_used": "best",
        "limit_used": 25,
        "total_comments_extracted": 2,
        "ai_analysis": {
            "category": "Question",
            "sentiment": "neutral",
            "summary": (
                "A Python developer is asking the community to recommend "
                "data analysis libraries. They highlight Pandas, Polars, and DuckDB "
                "as candidates. The post invites discussion of hidden gems."
            ),
            "key_quote": "What has your experience been? Any hidden gems worth trying?",
            "pain_points": ["Uncertainty about which library performs best at scale"],
            "mentioned_products": ["Pandas", "Polars", "DuckDB"],
            "solution_requests": ["Benchmark comparison tool for Python data libraries"],
            "willingness_to_pay": False,
        },
    }

    out = generate_docx([post_data, post_data], query="best python data analysis libraries")
    print(f"[OK] DOCX generated: {out}")
