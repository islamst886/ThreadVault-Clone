"""
youtube_docx_generator.py
--------------------------
Builds the YouTube Comment Archive DOCX report.

Structure:
  ┌─ Cover page (title, date, channels processed, totals)
  │
  └─ For each channel:
       Channel header (name, URL, subscribers, video count)
       └─ For each video (1-25):
            Video header (title, URL, views, likes, date, duration)
            └─ For each comment:
                 Comment block (rank, author, text, likes, date)
                 └─ For each reply:
                      Reply block (indented)
"""

import os
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm


# ── Colour palette ────────────────────────────────────────────────────────────
COL_YT_RED    = RGBColor(0xFF, 0x00, 0x00)
COL_DARK      = RGBColor(0x11, 0x11, 0x11)
COL_GREY      = RGBColor(0x66, 0x66, 0x66)
COL_LIGHT     = RGBColor(0xAA, 0xAA, 0xAA)
COL_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COL_ACCENT    = RGBColor(0x20, 0x6B, 0xC2)   # link blue


# ── Helpers ───────────────────────────────────────────────────────────────────

def _fmt_num(n: int) -> str:
    if n >= 1_000_000_000:
        return f"{n/1_000_000_000:.1f}B"
    if n >= 1_000_000:
        return f"{n/1_000_000:.1f}M"
    if n >= 1_000:
        return f"{n/1_000:.1f}K"
    return str(n)


def _set_run_color(run, color: RGBColor):
    run.font.color.rgb = color


def _set_para_spacing(para, before: int = 0, after: int = 0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def _add_hr(doc: Document):
    """Add a thin horizontal rule."""
    para = doc.add_paragraph()
    _set_para_spacing(para, before=40, after=40)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_hyperlink(para, text: str, url: str):
    """Add a clickable hyperlink run inside an existing paragraph."""
    part = para.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2060C0")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)


# ── Cover page ────────────────────────────────────────────────────────────────

def _build_cover(doc: Document, channels: list[dict], run_date: str):
    # Big title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(para, before=480, after=120)
    run = para.add_run("ThreadVault")
    run.font.size = Pt(36)
    run.bold = True
    _set_run_color(run, COL_DARK)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(para, before=0, after=360)
    run = para.add_run("YouTube Comment Archive")
    run.font.size = Pt(22)
    _set_run_color(run, COL_YT_RED)

    # Date
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(para, before=0, after=240)
    run = para.add_run(f"Generated: {run_date}")
    run.font.size = Pt(11)
    _set_run_color(run, COL_GREY)

    _add_hr(doc)

    # Channel summary
    total_videos   = sum(len(c.get("videos", [])) for c in channels)
    total_comments = sum(
        len(v.get("comments", []))
        for c in channels
        for v in c.get("videos", [])
    )
    total_replies  = sum(
        len(cm.get("replies", []))
        for c in channels
        for v in c.get("videos", [])
        for cm in v.get("comments", [])
    )

    stats = [
        ("Channels processed", str(len(channels))),
        ("Videos analysed", f"{total_videos:,}"),
        ("Top-level comments", f"{total_comments:,}"),
        ("Replies included", f"{total_replies:,}"),
    ]
    for label, value in stats:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(para, before=60, after=60)
        r1 = para.add_run(f"{label}: ")
        r1.font.size = Pt(11)
        _set_run_color(r1, COL_GREY)
        r2 = para.add_run(value)
        r2.font.size = Pt(11)
        r2.bold = True
        _set_run_color(r2, COL_DARK)

    _add_hr(doc)

    # Channel index
    para = doc.add_paragraph()
    _set_para_spacing(para, before=240, after=120)
    run = para.add_run("Channels in this report")
    run.font.size = Pt(14)
    run.bold = True
    _set_run_color(run, COL_DARK)

    for i, ch in enumerate(channels, 1):
        para = doc.add_paragraph()
        _set_para_spacing(para, before=40, after=40)
        r1 = para.add_run(f"  {i}. {ch.get('channel_name', 'Unknown Channel')}  ")
        r1.font.size = Pt(11)
        r1.bold = True
        _add_hyperlink(para, ch.get("channel_url", ""), ch.get("channel_url", ""))
        r2 = para.add_run(
            f"  •  {_fmt_num(ch.get('subscriber_count', 0))} subscribers"
            f"  •  {len(ch.get('videos', []))} videos"
        )
        r2.font.size = Pt(10)
        _set_run_color(r2, COL_GREY)

    doc.add_page_break()


# ── Channel section ───────────────────────────────────────────────────────────

def _build_channel_section(doc: Document, channel: dict, ch_idx: int):
    name    = channel.get("channel_name", "Unknown Channel")
    url     = channel.get("channel_url", "")
    subs    = channel.get("subscriber_count", 0)
    videos  = channel.get("videos", [])

    # Channel header banner
    para = doc.add_paragraph()
    _set_para_spacing(para, before=240, after=60)
    run = para.add_run(f"CHANNEL {ch_idx}: {name.upper()}")
    run.font.size = Pt(18)
    run.bold = True
    _set_run_color(run, COL_YT_RED)

    # Channel metadata row
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=40)
    run = para.add_run("URL: ")
    run.font.size = Pt(10)
    run.bold = True
    _set_run_color(run, COL_GREY)
    _add_hyperlink(para, url, url)

    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=40)
    meta = (
        f"Subscribers: {_fmt_num(subs)}"
        f"  ·  Videos analysed: {len(videos)}"
        f"  ·  Total channel views: {_fmt_num(channel.get('total_view_count', 0))}"
    )
    if channel.get("country"):
        meta += f"  ·  Country: {channel['country']}"
    run = para.add_run(meta)
    run.font.size = Pt(10)
    _set_run_color(run, COL_GREY)

    desc = channel.get("description", "").strip()
    if desc:
        para = doc.add_paragraph()
        _set_para_spacing(para, before=40, after=80)
        run = para.add_run(desc[:400] + ("…" if len(desc) > 400 else ""))
        run.font.size = Pt(9.5)
        run.italic = True
        _set_run_color(run, COL_GREY)

    _add_hr(doc)

    # Videos
    for v_idx, video in enumerate(videos, 1):
        _build_video_section(doc, video, v_idx, len(videos))

    doc.add_page_break()


# ── Video section ─────────────────────────────────────────────────────────────

def _build_video_section(
    doc: Document, video: dict, v_idx: int, total_videos: int
):
    title    = video.get("title", "Untitled Video")
    url      = video.get("url", "")
    views    = video.get("view_count", 0)
    likes    = video.get("like_count", 0)
    comments = video.get("comments", [])
    pub_date = video.get("published_at", "")[:10]
    duration = video.get("duration", "")

    # Video heading
    para = doc.add_paragraph()
    _set_para_spacing(para, before=200, after=40)
    run = para.add_run(f"VIDEO {v_idx} of {total_videos}  ·  ")
    run.font.size = Pt(10)
    run.bold = True
    _set_run_color(run, COL_LIGHT)
    run2 = para.add_run(title)
    run2.font.size = Pt(13)
    run2.bold = True
    _set_run_color(run2, COL_DARK)

    # Video URL
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=30)
    run = para.add_run("Link: ")
    run.font.size = Pt(9.5)
    run.bold = True
    _set_run_color(run, COL_GREY)
    _add_hyperlink(para, url, url)

    # Video stats
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=60)
    stats_text = (
        f"Views: {_fmt_num(views)}"
        f"  ·  Likes: {_fmt_num(likes)}"
        f"  ·  Comments: {_fmt_num(video.get('comment_count', len(comments)))}"
    )
    if pub_date:
        stats_text += f"  ·  Published: {pub_date}"
    if duration:
        stats_text += f"  ·  Duration: {duration}"
    run = para.add_run(stats_text)
    run.font.size = Pt(9.5)
    _set_run_color(run, COL_GREY)

    if not comments:
        para = doc.add_paragraph()
        _set_para_spacing(para, before=40, after=80)
        run = para.add_run("⚠ No comments available (comments may be disabled).")
        run.font.size = Pt(9.5)
        run.italic = True
        _set_run_color(run, COL_LIGHT)
        _add_hr(doc)
        return

    # Comments header
    para = doc.add_paragraph()
    _set_para_spacing(para, before=40, after=40)
    run = para.add_run(f"Top Comments ({len(comments):,} exported, sorted by relevance)")
    run.font.size = Pt(10)
    run.bold = True
    _set_run_color(run, COL_DARK)

    # Each comment
    for c_idx, comment in enumerate(comments, 1):
        _build_comment(doc, comment, c_idx, indent=False)

    _add_hr(doc)


# ── Comment block ─────────────────────────────────────────────────────────────

def _build_comment(doc: Document, comment: dict, rank: int, indent: bool):
    prefix = "    " if indent else ""
    author = comment.get("author", "Unknown")
    text   = comment.get("text", "")
    likes  = comment.get("likes", 0)
    date   = comment.get("published_at", "")

    # Author + meta line
    para = doc.add_paragraph()
    _set_para_spacing(para, before=80 if not indent else 40, after=20)
    if not indent:
        r = para.add_run(f"#{rank}  ")
        r.font.size = Pt(8)
        _set_run_color(r, COL_LIGHT)
    r1 = para.add_run(f"{prefix}@{author}")
    r1.font.size = Pt(9.5)
    r1.bold = True
    _set_run_color(r1, COL_ACCENT if not indent else COL_GREY)
    r2 = para.add_run(
        f"  ·  👍 {_fmt_num(likes)}"
        + (f"  ·  {date}" if date else "")
        + (f"  ·  {comment.get('reply_count', 0)} replies" if not indent and comment.get("reply_count", 0) else "")
    )
    r2.font.size = Pt(8.5)
    _set_run_color(r2, COL_LIGHT)

    # Comment text
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=40)
    # Indent replies visually
    if indent:
        para.paragraph_format.left_indent = Cm(1.5)
    else:
        para.paragraph_format.left_indent = Cm(0.5)

    run = para.add_run(text)
    run.font.size = Pt(9.5 if not indent else 9)
    _set_run_color(run, COL_DARK if not indent else COL_GREY)

    # Replies
    replies = comment.get("replies", [])
    if replies:
        for reply in replies:
            _build_comment(doc, reply, 0, indent=True)


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_youtube_docx(
    channels: list[dict],
    output_dir: str,
    progress_cb=None,
) -> str:
    """
    Build and save the YouTube Comment Archive DOCX.
    Returns the absolute path to the saved file.
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    # ── Page setup (A4, narrow margins) ──────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Default paragraph style ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    run_date = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # ── Cover ─────────────────────────────────────────────────────────────────
    _build_cover(doc, channels, run_date)

    # ── Per-channel sections ──────────────────────────────────────────────────
    for ch_idx, channel in enumerate(channels, 1):
        if progress_cb:
            progress_cb(ch_idx, len(channels), channel.get("channel_name", ""))
        _build_channel_section(doc, channel, ch_idx)

    # ── Save ──────────────────────────────────────────────────────────────────
    date_str       = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    channel_names  = "_".join(
        c.get("channel_name", "Channel")[:15].replace(" ", "") for c in channels[:3]
    )
    if len(channels) > 3:
        channel_names += f"_+{len(channels) - 3}more"
    filename = f"ThreadVault_YT_{date_str}_{channel_names}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return os.path.abspath(filepath)
