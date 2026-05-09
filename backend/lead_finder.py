"""
lead_finder.py
--------------
Backend engine for the Validation Hub → Find Leads feature.

Pipeline:
  1. Call Gemini API to generate 8 targeted Google search queries
     based on the user's SaaS description, target customer, and problems.
  2. Run each query through google_crawler.crawl_google_for_reddit_urls().
  3. Extract each Reddit post via reddit_extractor.extract_post().
  4. Run AI analysis via ai_analyzer.analyze_post().
  5. Filter: keep only Pain Point / Solution Request / Money Talk + negative/mixed.
  6. Extract post authors + high-scoring commenters as leads.
  7. Deduplicate by username, score 1–10, return sorted leads list.
"""

from __future__ import annotations

import asyncio
import json
import logging
import math
import os
import re
from dataclasses import dataclass, field
from typing import Any, Callable, Optional

logger = logging.getLogger(__name__)

# ── Lead job status constants ──────────────────────────────────────────────────
STATUS_QUEUED      = "queued"
STATUS_GENERATING  = "generating_queries"
STATUS_CRAWLING    = "crawling"
STATUS_EXTRACTING  = "extracting"
STATUS_ANALYZING   = "analyzing"
STATUS_SCORING     = "scoring"
STATUS_COMPLETE    = "complete"
STATUS_ERROR       = "error"


# ── In-memory lead job store (keyed by lead_job_id) ───────────────────────────
_lead_jobs: dict[str, "LeadJob"] = {}


def get_lead_jobs() -> dict[str, "LeadJob"]:
    return _lead_jobs


@dataclass
class LeadJob:
    """In-memory state for a single lead-finding run."""
    job_id:          str
    saas_description: str
    target_customer: str
    problems:        str
    subreddits:      str
    depth:           str           # "quick" | "standard" | "deep"
    status:          str = STATUS_QUEUED
    substatus:       Optional[str] = None
    error:           Optional[str] = None
    queries:         list[str] = field(default_factory=list)
    urls_found:      int = 0
    posts_analyzed:  int = 0
    leads:           list[dict] = field(default_factory=list)
    # Progress counters for the frontend
    total_queries:   int = 0
    queries_done:    int = 0


# ── Depth → pages mapping ──────────────────────────────────────────────────────
_DEPTH_PAGES = {
    "quick":    3,
    "standard": 10,
    "deep":     20,
}

# ── Category + sentiment filters ──────────────────────────────────────────────
_VALID_CATEGORIES = {"Pain Point", "Solution Request", "Money Talk"}
_VALID_SENTIMENTS = {"negative", "mixed"}


# ── Gemini: generate search queries ───────────────────────────────────────────

async def _generate_queries(
    saas_description: str,
    target_customer:  str,
    problems:         str,
) -> list[str]:
    """
    Asks Gemini to produce 8 site:reddit.com search queries.
    Returns the list; falls back to a simple default on any error.
    """
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key:
        logger.warning("[LeadFinder] GEMINI_API_KEY not set — using fallback queries.")
        return _fallback_queries(saas_description, target_customer, problems)

    prompt = (
        f"The user is validating a SaaS that does: {saas_description}\n"
        f"Target customer: {target_customer}\n"
        f"Problems it solves:\n{problems}\n\n"
        "Generate 8 Google search queries using site:reddit.com that would find "
        "Reddit posts where the target customers are complaining about these exact "
        "problems or asking for solutions.\n\n"
        "Return ONLY a JSON array of 8 query strings. Make them specific and "
        "conversational — the kind of language a frustrated user would actually "
        "use when posting on Reddit.\n\n"
        "Example style:\n"
        "'site:reddit.com personal trainer scheduling nightmare clients'\n"
        "'site:reddit.com how do personal trainers manage client payments'\n\n"
        "Return only the JSON array, nothing else."
    )

    try:
        from google import genai       # type: ignore
        from google.genai import types # type: ignore

        model = os.environ.get("GEMINI_MODEL", "gemini-3.1-flash-lite-preview")
        client = genai.Client(api_key=api_key)

        response = await asyncio.to_thread(
            client.models.generate_content,
            model=model,
            contents=prompt,
        )
        raw = response.text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```$", "", raw)
        queries = json.loads(raw)
        if isinstance(queries, list):
            return [str(q) for q in queries[:8]]
    except Exception as exc:
        logger.warning("[LeadFinder] Query generation failed: %s", exc)

    return _fallback_queries(saas_description, target_customer, problems)


def _fallback_queries(saas_description: str, target_customer: str, problems: str) -> list[str]:
    """Simple keyword-based fallback when Gemini is unavailable."""
    # Take first problem line as a short keyword phrase
    first_problem = problems.strip().split("\n")[0][:60] if problems.strip() else saas_description[:60]
    return [
        f"site:reddit.com {target_customer} {first_problem}",
        f"site:reddit.com {target_customer} struggling with management",
        f"site:reddit.com {target_customer} looking for software tool",
        f"site:reddit.com {target_customer} recommend app help",
        f"site:reddit.com {saas_description[:40]} problem",
        f"site:reddit.com {first_problem} solution",
        f"site:reddit.com {target_customer} workflow frustration",
        f"site:reddit.com {target_customer} manual process automation",
    ]


# ── Lead scoring ───────────────────────────────────────────────────────────────

def _score_lead(lead: dict) -> int:
    score = 0
    category = lead.get("category", "")
    sentiment = lead.get("sentiment", "")
    post_score = lead.get("post_score", 0) or 0
    appearances = lead.get("appearances", 1)
    comment_score = lead.get("comment_score", 0) or 0

    if category == "Solution Request":
        score += 3
    elif category == "Money Talk":
        score += 3
    elif category == "Pain Point" and sentiment in ("negative", "mixed"):
        score += 2

    if post_score > 10:
        score += 1
    if appearances > 1:
        score += 1
    if comment_score > 20:
        score += 1

    return min(score, 10)


def _priority_label(score: int) -> str:
    if score >= 8:
        return "HIGH PRIORITY"
    if score >= 5:
        return "GOOD LEAD"
    return "WEAK LEAD"


# ── Extract leads from a single post ──────────────────────────────────────────

def _extract_leads_from_post(post: dict) -> list[dict]:
    """
    Returns 0, 1, or more lead dicts extracted from a post + its comments.
    """
    leads: list[dict] = []
    ai = post.get("ai_analysis") or {}
    category  = ai.get("category", "")
    sentiment = ai.get("sentiment", "neutral")

    # Filter: only qualifying categories/sentiments become leads
    if category not in _VALID_CATEGORIES:
        return leads
    if sentiment not in _VALID_SENTIMENTS:
        return leads

    p        = post.get("post", {})
    author   = p.get("author", "")
    if not author or author in ("[deleted]", "[removed]", "AutoModerator"):
        author = None

    pain_points = ai.get("pain_points", [])
    pain_text   = " | ".join(pain_points) if pain_points else ""

    base = {
        "post_title":        p.get("title", ""),
        "post_url":          p.get("url", ""),
        "subreddit":         p.get("subreddit", ""),
        "post_date":         p.get("posted_at", ""),
        "pain_point_summary": pain_text,
        "key_quote":         ai.get("key_quote", ""),
        "category":          category,
        "sentiment":         sentiment,
        "post_score":        p.get("score", 0),
        "comment_score":     0,
        "appearances":       1,
        "status":            "new",
    }

    # Post author as a lead
    if author:
        lead = {**base, "reddit_username": author, "lead_type": "poster"}
        lead["score"] = _score_lead(lead)
        lead["priority"] = _priority_label(lead["score"])
        leads.append(lead)

    # High-scoring commenters as leads
    for comment in post.get("comments", []):
        if comment.get("depth", 0) != 0:
            continue
        c_score  = comment.get("score", 0) or 0
        c_author = comment.get("author", "")
        if c_score < 5:
            continue
        if not c_author or c_author in ("[deleted]", "[removed]", "AutoModerator"):
            continue
        c_body = comment.get("body", "").lower()
        # Only add commenter if their comment contains a pain/frustration signal
        pain_signals = ["struggle", "hard", "difficult", "problem", "hate", "wish",
                        "can't", "cannot", "annoying", "frustrat", "overwhelm",
                        "nightmare", "manual", "tedious", "time-consuming"]
        if not any(sig in c_body for sig in pain_signals):
            continue
        lead = {
            **base,
            "reddit_username": c_author,
            "lead_type":       "commenter",
            "comment_score":   c_score,
            "key_quote":       comment.get("body", "")[:200],
        }
        lead["score"] = _score_lead(lead)
        lead["priority"] = _priority_label(lead["score"])
        leads.append(lead)

    return leads


# ── Deduplication & merge ──────────────────────────────────────────────────────

def _merge_leads(raw_leads: list[dict]) -> list[dict]:
    """
    Deduplicate by reddit_username. If same user appears multiple times,
    merge their entries; bump appearances count → higher priority signal.
    """
    merged: dict[str, dict] = {}
    for lead in raw_leads:
        username = lead["reddit_username"].lower()
        if username not in merged:
            merged[username] = dict(lead)
        else:
            existing = merged[username]
            existing["appearances"] = existing.get("appearances", 1) + 1
            # Keep highest post_score and comment_score
            existing["post_score"]    = max(existing.get("post_score", 0),    lead.get("post_score", 0))
            existing["comment_score"] = max(existing.get("comment_score", 0), lead.get("comment_score", 0))
            # Accumulate pain_point_summary
            ep = existing.get("pain_point_summary", "")
            np = lead.get("pain_point_summary", "")
            if np and np not in ep:
                existing["pain_point_summary"] = f"{ep} | {np}".strip(" |")
            # Escalate category priority
            cat_priority = {"Money Talk": 3, "Solution Request": 2, "Pain Point": 1}
            if cat_priority.get(lead["category"], 0) > cat_priority.get(existing["category"], 0):
                existing["category"] = lead["category"]

    # Recompute score now that appearances are merged
    result = []
    for lead in merged.values():
        lead["score"]    = _score_lead(lead)
        lead["priority"] = _priority_label(lead["score"])
        result.append(lead)

    result.sort(key=lambda l: l["score"], reverse=True)
    return result


# ── Main pipeline (called from FastAPI background task) ───────────────────────

async def run_lead_finder(
    job_id: str,
    status_cb: Optional[Callable[[str, Optional[str]], None]] = None,
) -> None:
    """
    Full lead-finding pipeline. Mutates the LeadJob in _lead_jobs.

    Args:
        job_id:    Key in _lead_jobs dict.
        status_cb: Optional callback(status, substatus) for live UI updates.
    """
    # type: ignore suppresses Pyright/VS Code static analysis false positives.
    # At runtime, 'backend/' is the working directory so these import correctly.
    from google_crawler import crawl_multiple_google_queries_sync  # type: ignore
    import reddit_extractor as _re  # type: ignore
    import ai_analyzer      as _ai  # type: ignore

    job = _lead_jobs.get(job_id)
    if not job:
        logger.error("[LeadFinder] Job %s not found.", job_id)
        return

    def _set(status: str, sub: Optional[str] = None):
        job.status    = status
        job.substatus = sub
        if status_cb:
            status_cb(status, sub)

    try:
        # ── Step 1: Generate queries ─────────────────────────────────────────
        _set(STATUS_GENERATING, "Asking Gemini to generate search queries…")
        queries = await _generate_queries(
            job.saas_description,
            job.target_customer,
            job.problems,
        )
        job.queries       = queries
        job.total_queries = len(queries)
        logger.info("[LeadFinder] %d queries generated for job %s", len(queries), job_id)

        # ── Step 2: Crawl all queries sequentially in one session ────────────
        max_pages = _DEPTH_PAGES.get(job.depth, 10)
        all_urls: set[str] = set()

        def _crawl_status(idx: int, q: str):
            _set(STATUS_CRAWLING, f"Searching Google: query {idx}/{len(queries)} — {q[:60]}…")
            job.queries_done = idx

        try:
            urls = await asyncio.to_thread(
                crawl_multiple_google_queries_sync,
                queries,
                max_pages=max_pages,
                headless=False,
                status_cb=_crawl_status,
            )
            all_urls.update(urls)
            logger.info("[LeadFinder] Bulk crawl complete, returned %d total URLs", len(urls))
        except Exception as exc:
            logger.warning("[LeadFinder] Bulk crawl failed: %s", exc)
        
        job.queries_done = len(queries)

        job.urls_found = len(all_urls)
        if not all_urls:
            job.status = STATUS_COMPLETE
            job.substatus = None
            job.leads = []
            logger.warning("[LeadFinder] No URLs found for job %s.", job_id)
            return

        # ── Step 3: Extract + analyze posts ──────────────────────────────────
        url_list = list(all_urls)
        raw_leads: list[dict] = []

        for i, url in enumerate(url_list, 1):
            _set(STATUS_EXTRACTING, f"Extracting post {i}/{len(url_list)}…")
            try:
                post = await _re.extract_post(url, sort="best", limit=25)
            except Exception as exc:
                logger.warning("[LeadFinder] Extract failed for %s: %s", url, exc)
                continue

            if not post:
                continue

            _set(STATUS_ANALYZING, f"Analyzing post {i}/{len(url_list)} with AI…")
            try:
                await _ai.analyze_post(post)
            except Exception as exc:
                logger.warning("[LeadFinder] AI analysis failed for %s: %s", url, exc)

            # ── Step 4-6: Extract & collect leads from this post ─────────────
            post_leads = _extract_leads_from_post(post)
            raw_leads.extend(post_leads)
            job.posts_analyzed = i

        # ── Step 7: Deduplicate and score ─────────────────────────────────────
        _set(STATUS_SCORING, "Deduplicating and scoring leads…")
        job.leads = _merge_leads(raw_leads)
        logger.info("[LeadFinder] Job %s complete — %d leads found.", job_id, len(job.leads))

        job.status    = STATUS_COMPLETE
        job.substatus = None

    except Exception as exc:
        logger.error("[LeadFinder] Job %s crashed: %s", job_id, exc, exc_info=True)
        job.status = STATUS_ERROR
        job.error  = str(exc)
        job.substatus = None


# ── Message generation ─────────────────────────────────────────────────────────

_STYLE_PROMPTS = {
    "direct": (
        "Direct and honest. Get straight to the point. "
        "Mention exactly what they said and ask one direct validation question."
    ),
    "curious": (
        "Curious and exploratory. Sound genuinely fascinated by their problem. "
        "Frame yourself as a researcher, not a founder. "
        "Ask an open-ended question about how they currently deal with it."
    ),
    "story": (
        "Story-led. Start with a brief personal anecdote (1 sentence) about why "
        "you noticed their problem. Make it feel like you're sharing something in common. "
        "End with a question about whether they've found anything that helps."
    ),
}

_MSG_SYSTEM_PROMPT = (
    "You are an expert at writing genuine, non-spammy cold outreach messages "
    "for SaaS market validation on Reddit. Your messages sound like a real founder "
    "reaching out for honest feedback, NOT a salesperson. They are short, specific, "
    "and reference something the recipient actually said."
)

_MSG_USER_TEMPLATE = """\
Write a Reddit DM for SaaS market validation.

My SaaS: {saas_description}

About this lead:
  Reddit username: {username}
  What they posted about: {post_title}
  Their key quote: {key_quote}
  Their pain point: {pain_point}
  Subreddit: r/{subreddit}
  Post date: {post_date}
  Category: {category}

Message style: {style_instruction}

Rules for the message:
1. Start by referencing something SPECIFIC they said — never a generic opener
2. Maximum 4 sentences total
3. One specific validation question at the end — not "can I show you my product"
   but a genuine question like "is this something you'd actually pay for?"
4. No links, no pitch, no CTA to sign up
5. Sound like a curious founder, not a salesperson
6. Do NOT mention the product name or price

Return ONLY a JSON object with exactly these fields:
{{
  "subject": "DM subject line (5 words max)",
  "message": "the full message text",
  "personalization_note": "one sentence explaining what specific thing you referenced from their post",
  "validation_question": "the question you ended with",
  "estimated_read_time": "X seconds"
}}"""


async def _call_gemini_for_message(
    prompt: str,
    style_key: str,
) -> dict:
    """Single Gemini call for one message style. Returns parsed dict."""
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    model   = os.environ.get("GEMINI_MODEL", "gemini-3.1-flash-lite-preview")

    if not api_key:
        return _fallback_message(style_key)

    try:
        from google import genai       # type: ignore
        from google.genai import types # type: ignore

        client = genai.Client(api_key=api_key)
        response = await asyncio.to_thread(
            client.models.generate_content,
            model=model,
            contents=prompt,
            config=types.GenerateContentConfig(
                system_instruction=_MSG_SYSTEM_PROMPT,
                response_mime_type="application/json",
            ),
        )
        raw = response.text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```$", "", raw)
        data = json.loads(raw)
        data["style"] = style_key
        return data
    except Exception as exc:
        logger.warning("[LeadFinder] Message gen failed (%s): %s", style_key, exc)
        return _fallback_message(style_key)


def _fallback_message(style_key: str) -> dict:
    return {
        "style":                style_key,
        "subject":              "Quick question from a founder",
        "message":              (
            "Hey, I came across your post and it really resonated with me. "
            "I'm building something to help with exactly that problem. "
            "Would you be open to telling me more about how you currently deal with it?"
        ),
        "personalization_note": "Generic fallback — Gemini API unavailable.",
        "validation_question":  "Would you be open to telling me more about how you currently deal with it?",
        "estimated_read_time":  "10 seconds",
    }


async def generate_outreach_message(
    lead: dict,
    saas_description: str,
) -> dict:
    """
    Generates 3 message versions (direct / curious / story) in parallel.

    Returns:
        {
          "versions": [
            {"style": "direct",  "subject": ..., "message": ..., ...},
            {"style": "curious", ...},
            {"style": "story",   ...}
          ]
        }
    """
    prompts = {}
    for style_key, style_instr in _STYLE_PROMPTS.items():
        prompts[style_key] = _MSG_USER_TEMPLATE.format(
            saas_description = saas_description[:300],
            username         = lead.get("reddit_username", ""),
            post_title       = lead.get("post_title", "")[:200],
            key_quote        = lead.get("key_quote", "")[:300],
            pain_point       = lead.get("pain_point_summary", "")[:300],
            subreddit        = lead.get("subreddit", ""),
            post_date        = (lead.get("post_date") or "")[:10],
            category         = lead.get("category", ""),
            style_instruction = style_instr,
        )

    results = await asyncio.gather(
        _call_gemini_for_message(prompts["direct"],  "direct"),
        _call_gemini_for_message(prompts["curious"], "curious"),
        _call_gemini_for_message(prompts["story"],   "story"),
    )
    return {"versions": list(results)}


async def generate_messages_batch(
    leads: list[dict],
    saas_description: str,
    max_concurrent: int = 3,
) -> list[dict]:
    """
    Generates messages for a list of leads with bounded concurrency.
    Returns list of {lead, versions} dicts.
    """
    semaphore = asyncio.Semaphore(max_concurrent)

    async def _gen_one(lead: dict) -> dict:
        async with semaphore:
            result = await generate_outreach_message(lead, saas_description)
            return {"lead": lead, "versions": result["versions"]}

    return list(await asyncio.gather(*[_gen_one(l) for l in leads]))


def generate_messages_docx(
    batch_results: list[dict],
    output_path: str,
) -> str:
    """
    Creates a DOCX with one page per lead message set.
    Returns the output_path.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Cover page
    title_para = doc.add_heading("ThreadVault — Outreach Messages", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Generated {len(batch_results)} personalized message sets. "
        "Each message references something specific that person said."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    style_labels = {"direct": "Version A: Direct", "curious": "Version B: Curious", "story": "Version C: Story-led"}

    for entry in batch_results:
        lead     = entry.get("lead", {})
        versions = entry.get("versions", [])

        username  = lead.get("reddit_username", "unknown")
        subreddit = lead.get("subreddit", "")
        key_quote = lead.get("key_quote", "")

        doc.add_heading(f"u/{username}  •  r/{subreddit}", level=1)

        if key_quote:
            p = doc.add_paragraph()
            p.add_run("What they said: ").bold = True
            p.add_run(f'"{key_quote[:250]}"').italic = True

        doc.add_paragraph(f"Pain point: {lead.get('pain_point_summary', 'N/A')[:200]}")
        doc.add_paragraph(f"Score: {lead.get('score', '?')}/10  |  Category: {lead.get('category', '?')}")
        doc.add_paragraph("")

        for v in versions:
            style_key = v.get("style", "")
            label     = style_labels.get(style_key, style_key.title())

            doc.add_heading(label, level=2)

            subj_para = doc.add_paragraph()
            subj_para.add_run("Subject: ").bold = True
            subj_para.add_run(v.get("subject", ""))

            msg_para = doc.add_paragraph(v.get("message", ""))

            note_para = doc.add_paragraph()
            note_para.add_run("ℹ Personalized by: ").bold = True
            note_para.add_run(v.get("personalization_note", ""))
            note_run = note_para.runs[-1]
            note_run.font.color.rgb = RGBColor(0x8A, 0x95, 0xB0)
            note_run.font.size      = Pt(9)

            doc.add_paragraph("")

        doc.add_page_break()

    doc.save(output_path)
    return output_path


# ── Pipeline helpers ───────────────────────────────────────────────────────────

async def generate_followup_message(lead: dict, saas_description: str) -> str:
    """
    Generates a single follow-up DM sentence for a lead who hasn't replied.
    """
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    model   = os.environ.get("GEMINI_MODEL", "gemini-3.1-flash-lite-preview")

    key_quote = lead.get("key_quote") or lead.get("pain_point_summary") or "that problem you mentioned"

    if not api_key:
        return (
            f"Hey, just checking back — did you ever find anything that helps with "
            f"that challenge you mentioned?"
        )

    prompt = (
        f"Write a 1-sentence follow-up DM for someone who didn't reply to my previous message "
        f"about {saas_description}.\n\n"
        f"Their original pain point: {key_quote[:200]}\n\n"
        f"Rules:\n"
        f"- Maximum 1 sentence\n"
        f"- No pressure, no pitch\n"
        f"- Reference their original pain point briefly\n"
        f"- End with a yes/no question\n"
        f"Return only the message text, nothing else."
    )

    try:
        from google import genai       # type: ignore
        client = genai.Client(api_key=api_key)
        response = await asyncio.to_thread(
            client.models.generate_content,
            model=model,
            contents=prompt,
        )
        return response.text.strip().strip('"\'')
    except Exception as exc:
        logger.warning("[LeadFinder] Follow-up gen failed: %s", exc)
        return "Hey, just checking back — did you ever find a solution for that challenge you mentioned?"


def export_pipeline_docx(
    stages:           dict,          # {stage_key: [lead_dict, ...]}
    profile_name:     str,
    saas_description: str,
    target_customer:  str,
    output_path:      str,
) -> str:
    """
    Exports the full pipeline as a structured market-validation research report DOCX.
    """
    from collections import Counter
    from datetime import datetime

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Cover page ───────────────────────────────────────────────────────────
    title = doc.add_heading(f"Market Validation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(profile_name)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"SaaS description: {saas_description}")
    doc.add_paragraph(f"Target customer: {target_customer}")
    doc.add_page_break()

    # ── Stats ────────────────────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)

    all_leads   = [l for s in stages.values() for l in s]
    total       = len(all_leads)
    contacted_n = sum(len(stages.get(s, [])) for s in ("contacted", "followup", "replied", "validated"))
    replied_n   = sum(len(stages.get(s, [])) for s in ("replied", "validated"))
    validated_n = len(stages.get("validated", []))
    reply_pct   = f"{round(replied_n / contacted_n * 100)}%" if contacted_n else "N/A"

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Light Grid"
    for i, (label, val) in enumerate([
        ("Total Leads Found",       total),
        ("Contacted",               contacted_n),
        ("Replies Received",        replied_n),
        ("Reply Rate",              reply_pct),
        ("Validated Conversations", validated_n),
    ]):
        row         = tbl.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(val)

    doc.add_paragraph(
        f"\nIndustry benchmark for cold outreach on Reddit: 3–10% reply rate. "
        f"Your current rate: {reply_pct}."
    )
    doc.add_page_break()

    # ── Validated conversations ───────────────────────────────────────────────
    validated_leads = stages.get("validated", [])
    if validated_leads:
        doc.add_heading("Validated Conversations", level=1)
        doc.add_paragraph(
            "These leads confirmed the problem exists and had a meaningful conversation:"
        )
        for lead in validated_leads:
            doc.add_heading(
                f"u/{lead.get('reddit_username','?')}  •  r/{lead.get('subreddit', '?')}",
                level=2,
            )
            if lead.get("key_quote"):
                p = doc.add_paragraph()
                p.add_run("Key Quote: ").bold = True
                p.add_run(f"\"{lead['key_quote'][:300]}\"").italic = True
            p = doc.add_paragraph()
            p.add_run("Pain Point: ").bold = True
            p.add_run(lead.get("pain_point_summary", "N/A")[:300])
            if lead.get("notes"):
                p = doc.add_paragraph()
                p.add_run("Notes: ").bold = True
                p.add_run(lead["notes"])
            doc.add_paragraph("")
        doc.add_page_break()

    # ── Pain-point frequency (Market Evidence) ────────────────────────────────
    doc.add_heading("Market Evidence — Pain Points by Frequency", level=1)
    doc.add_paragraph(
        "The following pain points were mentioned by leads discovered through Reddit research. "
        "This is your validation evidence to share with investors or co-founders."
    )
    doc.add_paragraph("")

    pain_counter: Counter = Counter()
    for lead in all_leads:
        pain = lead.get("pain_point_summary", "")
        for item in pain.split(" | "):
            item = item.strip()
            if len(item) > 5:
                pain_counter[item] += 1

    if pain_counter:
        for pain, count in pain_counter.most_common(25):
            p = doc.add_paragraph()
            run = p.add_run(f"• {pain}")
            count_run = p.add_run(f"  — mentioned by {count} lead{'s' if count > 1 else ''}")
            count_run.font.color.rgb = RGBColor(0x8A, 0x95, 0xB0)
            count_run.font.size      = Pt(9)
    else:
        doc.add_paragraph("No pain points recorded yet.")

    doc.add_page_break()

    # ── All leads by stage ────────────────────────────────────────────────────
    stage_labels = {
        "new":          "New Leads",
        "contacted":    "Contacted",
        "followup":     "Follow-up Due",
        "replied":      "Replied",
        "validated":    "Validated",
        "rejected":     "Not Interested",
    }
    doc.add_heading("All Leads by Stage", level=1)
    for stage_key, label in stage_labels.items():
        s_leads = stages.get(stage_key, [])
        if not s_leads:
            continue
        doc.add_heading(f"{label} ({len(s_leads)})", level=2)
        for lead in s_leads:
            p = doc.add_paragraph()
            p.add_run(f"u/{lead.get('reddit_username','?')}").bold = True
            p.add_run(f"  r/{lead.get('subreddit','?')}  Score: {lead.get('score','?')}/10")
            if lead.get("key_quote"):
                doc.add_paragraph(f"  \"{lead['key_quote'][:150]}\"").italic = True
            if lead.get("notes"):
                p2 = doc.add_paragraph()
                p2.add_run("  Notes: ").bold = True
                p2.add_run(lead["notes"])

    doc.save(output_path)
    return output_path
