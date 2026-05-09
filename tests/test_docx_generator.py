"""
tests/test_docx_generator.py
-----------------------------
Unit tests for docx_generator.py.

These tests verify document structure, file creation, and formatting logic
using synthetic PostData / CommentData objects — no API calls required.
"""

import os
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

from docx_generator import (  # type: ignore
    generate_docx,
    _apply_run_style,
    _set_paragraph_indent,
    INDENT_PER_DEPTH,
)
from docx import Document
from docx.shared import Inches, Pt, RGBColor


# ── Fixtures ──────────────────────────────────────────────────────────────────

def _make_comment(depth=0, body="Test comment body", score=10, author="user1",
                  replies=None):
    return {
        "author": author,
        "body": body,
        "score": score,
        "posted_at": "2024-01-01 00:00 UTC",
        "depth": depth,
        "replies": replies or [],
        "media_note": None,
    }

def _make_post(title="Test Post", body="Test body", media_type="text",
               media_note=None, comments=None):
    post = {
        "url": "https://www.reddit.com/r/test/comments/abc123/test_post/",
        "title": title,
        "subreddit": "test",
        "author": "poster",
        "score": 500,
        "upvote_ratio": "95%",
        "num_comments": 20,
        "flair": None,
        "posted_at": "2024-01-01 00:00 UTC",
        "body": body,
        "media_note": media_note,
    }
    return {
        "post": post,
        "comments": comments or [],
        "sort_used": "best",
        "limit_used": 25,
        "total_comments_extracted": len(comments or []),
    }


# ── generate_docx — file creation ─────────────────────────────────────────────

class TestGenerateDocxFileCreation:
    def test_creates_file(self, tmp_path):
        post = _make_post()
        path = generate_docx([post], query="test query", output_dir=str(tmp_path))
        assert os.path.isfile(path)

    def test_filename_contains_query_slug(self, tmp_path):
        # Filename format: <query_slug>_<N>posts_<timestamp>.docx
        post = _make_post()
        path = generate_docx([post], query="test query", output_dir=str(tmp_path))
        fname = os.path.basename(path)
        assert fname.startswith("test_query_")
        assert "1posts" in fname
        assert fname.endswith(".docx")

    def test_filename_ends_with_docx(self, tmp_path):
        post = _make_post()
        path = generate_docx([post], query="test query", output_dir=str(tmp_path))
        assert path.endswith(".docx")

    def test_creates_output_dir_if_missing(self, tmp_path):
        nested = str(tmp_path / "newdir" / "subdir")
        post = _make_post()
        path = generate_docx([post], query="test query", output_dir=nested)
        assert os.path.isdir(nested)
        assert os.path.isfile(path)

    def test_raises_on_empty_posts(self, tmp_path):
        with pytest.raises(ValueError, match="No posts provided"):
            generate_docx([], query="test", output_dir=str(tmp_path))

    def test_multiple_posts_single_file(self, tmp_path):
        posts = [_make_post(f"Post {i}") for i in range(5)]
        path = generate_docx(posts, query="multi test", output_dir=str(tmp_path))
        # Only one file should be created per call
        files = list(tmp_path.glob("*.docx"))
        assert len(files) == 1

    def test_returns_absolute_path(self, tmp_path):
        post = _make_post()
        path = generate_docx([post], query="test", output_dir=str(tmp_path))
        assert os.path.isabs(path)


# ── generate_docx — document content ──────────────────────────────────────────

class TestGenerateDocxContent:
    def _load_doc(self, tmp_path, posts, query="test query"):
        path = generate_docx(posts, query=query, output_dir=str(tmp_path))
        return Document(path)

    def test_document_has_paragraphs(self, tmp_path):
        doc = self._load_doc(tmp_path, [_make_post()])
        assert len(doc.paragraphs) > 0

    def test_document_title_contains_query(self, tmp_path):
        doc = self._load_doc(tmp_path, [_make_post()], query="my search term")
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "my search term" in full_text

    def test_document_title_contains_threadVault(self, tmp_path):
        doc = self._load_doc(tmp_path, [_make_post()])
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "ThreadVault" in full_text

    def test_post_title_in_document(self, tmp_path):
        doc = self._load_doc(tmp_path, [_make_post(title="My Amazing Title")])
        texts = [p.text for p in doc.paragraphs]
        assert any("My Amazing Title" in t for t in texts)

    def test_post_body_in_document(self, tmp_path):
        doc = self._load_doc(tmp_path, [_make_post(body="Unique body content 12345")])
        texts = [p.text for p in doc.paragraphs]
        assert any("Unique body content 12345" in t for t in texts)

    def test_subreddit_in_metadata(self, tmp_path):
        doc = self._load_doc(tmp_path, [_make_post()])
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "r/test" in full_text

    def test_media_note_in_document(self, tmp_path):
        post = _make_post(
            media_note="[POST CONTAINS: image — not copied]",
            body="",
        )
        doc = self._load_doc(tmp_path, [post])
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "not copied" in full_text

    def test_comment_author_in_document(self, tmp_path):
        comment = _make_comment(author="famous_redditor", body="Great post!")
        post = _make_post(comments=[comment])
        doc = self._load_doc(tmp_path, [post])
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "famous_redditor" in full_text

    def test_comment_body_in_document(self, tmp_path):
        comment = _make_comment(body="This is a very specific comment text.")
        post = _make_post(comments=[comment])
        doc = self._load_doc(tmp_path, [post])
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "very specific comment text" in full_text

    def test_reply_in_document(self, tmp_path):
        reply = _make_comment(depth=1, author="replier", body="I replied here.")
        comment = _make_comment(replies=[reply])
        post = _make_post(comments=[comment])
        doc = self._load_doc(tmp_path, [post])
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "replier" in full_text
        assert "I replied here" in full_text

    def test_post_separator_present(self, tmp_path):
        posts = [_make_post("P1"), _make_post("P2")]
        doc = self._load_doc(tmp_path, posts)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "POST 1 of 2" in full_text
        assert "POST 2 of 2" in full_text

    def test_no_comments_message(self, tmp_path):
        post = _make_post(comments=[])
        doc = self._load_doc(tmp_path, [post])
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "No comments extracted" in full_text

    def test_flair_in_metadata(self, tmp_path):
        post = _make_post()
        post["post"]["flair"] = "Discussion"
        doc = self._load_doc(tmp_path, [post])
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Discussion" in full_text

    def test_removed_body_shown(self, tmp_path):
        post = _make_post(body="[removed]")
        doc = self._load_doc(tmp_path, [post])
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "[removed]" in full_text


# ── Helper function unit tests ─────────────────────────────────────────────────

class TestHelpers:
    def test_set_paragraph_indent_depth_0(self):
        doc = Document()
        p = doc.add_paragraph()
        _set_paragraph_indent(p, depth=0)
        assert p.paragraph_format.left_indent == Inches(0)

    def test_set_paragraph_indent_depth_1(self):
        doc = Document()
        p = doc.add_paragraph()
        _set_paragraph_indent(p, depth=1)
        assert p.paragraph_format.left_indent == Inches(INDENT_PER_DEPTH * 1)

    def test_set_paragraph_indent_depth_3(self):
        doc = Document()
        p = doc.add_paragraph()
        _set_paragraph_indent(p, depth=3)
        assert p.paragraph_format.left_indent == Inches(INDENT_PER_DEPTH * 3)

    def test_apply_run_style_bold(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        _apply_run_style(run, Pt(12), bold=True)
        assert run.font.bold is True

    def test_apply_run_style_italic(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        _apply_run_style(run, Pt(12), italic=True)
        assert run.font.italic is True

    def test_apply_run_style_size(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        _apply_run_style(run, Pt(14))
        assert run.font.size == Pt(14)

    def test_apply_run_style_color(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        color = RGBColor(0xFF, 0x00, 0x00)
        _apply_run_style(run, Pt(12), color=color)
        assert run.font.color.rgb == color
